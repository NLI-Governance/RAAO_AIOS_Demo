import os
from docx import Document
from langchain_community.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.schema import Document as LC_Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

def extract_text_with_metadata(docx_path):
    document = Document(docx_path)
    chunks = []
    current_section = "Untitled"
    current_chunk = ""
    line_counter = 0
    start_line = 1

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if para.style.name.startswith("Heading"):
            if current_chunk:
                chunks.append((current_chunk.strip(), current_section, start_line, line_counter))
                start_line = line_counter + 1
                current_chunk = ""
            current_section = text
        else:
            current_chunk += " " + text
        line_counter += 1

    if current_chunk:
        chunks.append((current_chunk.strip(), current_section, start_line, line_counter))

    return chunks

def build_faiss_index():
    input_dir = os.path.expanduser("~/Dropbox/EI_Cloud/ABL_Rev2_6_1_25/data/policies/RAAO_sources/")
    output_dir = os.path.expanduser("~/Dropbox/EI_Cloud/ABL_Rev2_6_1_25/data/policies/faiss_index/")
    os.makedirs(output_dir, exist_ok=True)

    splitter = RecursiveCharacterTextSplitter(chunk_size=600, chunk_overlap=50)
    embedding = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    all_docs = []

    for filename in os.listdir(input_dir):
        if filename.lower().endswith(".docx"):
            full_path = os.path.join(input_dir, filename)
            base_name = os.path.splitext(filename)[0]
            print(f"Processing: {filename}")

            extracted_chunks = extract_text_with_metadata(full_path)
            for chunk, section, start_line, end_line in extracted_chunks:
                for subchunk in splitter.split_text(chunk):
                    doc = LC_Document(
                        page_content=subchunk,
                        metadata={
                            "source": base_name,
                            "section": section,
                            "lines": f"L{start_line}-L{end_line}"
                        }
                    )
                    all_docs.append(doc)

    print(f"Embedding {len(all_docs)} chunks using MiniLM...")
    db = FAISS.from_documents(all_docs, embedding)
    db.save_local(output_dir)
    print(f"✅ Local FAISS index saved to: {output_dir}")

if __name__ == "__main__":
    build_faiss_index()
